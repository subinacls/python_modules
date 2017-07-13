#!/usr/bin/python
#
# Required:
#  pip install python-docx
#  pip install mistune

import sys
import os
import argparse
import base64
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from HTMLParser import HTMLParser
import mistune


class DocxHTMLParser(HTMLParser):
    def __init__(self, doc):
        self.tag = [None]
        self.doc = doc
        self.ctx = None
        HTMLParser.__init__(self)
    def handle_starttag(self, tag, attrs):
        #print "start tag:", tag
        self.tag.append(tag)
        if tag == 'p':
            self.ctx = self.doc.add_paragraph()
    def handle_endtag(self, tag):
        #print "end tag :", tag
        self.tag.pop()
        if tag == 'p':
            # TODO - should make ctx dict for nested <p> tags
            self.ctx = None
    def handle_data(self, data):
        tag = self.tag[-1] # peek at tag context
        #print "({}) : {}".format( tag, data )
        if tag == 'p':
            if self.ctx is None:
                self.ctx = self.doc.add_paragraph()
                #print "??????WHAT????? p tag data and null ctx... this shouldnt happen"
            if self.tag[-2] == 'li':
                if self.tag[-3] == 'ul':
                    self.doc.add_paragraph(data, style='ListBullet')
                elif self.tag[-3] == 'ol':
                    self.doc.add_paragraph(data, style='ListNumber')
            else:
                self.ctx.add_run(data)
        elif tag == 'a':
            # TODO - proper links
            if self.ctx is None:
                self.doc.add_paragraph(data)
            else:
                self.ctx.add_run(data)
        elif (tag == 'strong') or (tag == 'b'):
            if self.ctx is not None:
                self.ctx.add_run(data).bold = True
            else:
                pg = self.doc.add_paragraph()
                pg.add_run(data).bold = True
        elif (tag == 'em') or (tag == 'i'):
            if self.ctx is not None:
                self.ctx.add_run(data).italic = True
            else:
                pg = self.doc.add_paragraph()
                pg.add_run(data).italic = True
        elif tag == 'li':
            # get parent
            # mistune is wrapping li data in p tags
            #print data
            if data is not None and len(data) > 1:
                if self.tag[-2] == 'ul':
                    self.doc.add_paragraph(data, style='ListBullet')
                elif self.tag[-2] == 'ol':
                    self.doc.add_paragraph(data, style='ListNumber')
        elif tag is None:
            pass
        elif tag == 'ul' or tag == 'ol':
            pass # it's ok
        else:
            print "[!] WARNING: No handler for <{}> tag. Skipping...".format(tag)

def script_path():
    return os.path.dirname(os.path.realpath(__file__))


def parse( args ):
    findings = {}
    #common = {}
    # Import common findings
    #print "[*] Loading common findings..."
    #tree = ET.parse( open(script_path() + '/common.xml','r') )
    #root = tree.getroot()
    # Parse files
    for infile in args.inputfiles:
        print "[*] Reading " + infile.name
        tree = ET.parse( infile )
        root = tree.getroot().find('Report')
        print "[*] Parsing findings."
        for host in root.findall('ReportHost'):
            for item in host.findall('ReportItem'):
                severity = item.find("risk_factor").text
                plugin_id = item.get('pluginID')
                if severity != "None":
                    if plugin_id not in findings:
                        findings[plugin_id] = {}
                        findings[plugin_id]['hosts'] = []
                        try:
                            findings[plugin_id]['name'] = item.get('pluginName')
                            findings[plugin_id]['severity'] = severity
                            if item.find('cvss3_base_score') is not None:
                                findings[plugin_id]['cvss3_base_score'] = item.find('cvss3_base_score').text
                            if item.find('cvss3_vector') is not None:
                                findings[plugin_id]['cvss3_vector'] = item.find('cvss3_vector').text
                            findings[plugin_id]['description'] = item.find('description').text
                            findings[plugin_id]['recommendation'] = item.find('solution').text
                            if item.find('see_also') is not None:
                                findings[plugin_id]['resources'] = item.find('see_also').text
                                #print findings[plugin_id]['resources']
                            if item.findall('xref') is not None:
                                findings[plugin_id]['bugid'] = []
                                for xref in item.iter('xref'):
                                    findings[plugin_id]['bugid'].append(xref.text)
                            findings[plugin_id]['output'] = item.find('plugin_output').text
                        except Exception as e:
                            pass
                    findings[plugin_id]['hosts'].append(host.get('name'))

    print "[*] Generating report."
    document = Document(docx=script_path() + "/template.docx")
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    printFindings(document, findings, "Critical")
    printFindings(document, findings, "High")
    printFindings(document, findings, "Medium")
    printFindings(document, findings, "Low")
    document.save('report.docx')
    print "[*] Output saved to report.docx"

def printFindings(document, findings, severity):
    print "[*] {} Severity Findings:".format(severity)
    count = 0
    parser = DocxHTMLParser(document)
    alpha = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
    for c in range(0, len(alpha)): # poor man sort
        for f in findings:
            if findings[f]['severity'] == severity and findings[f]['name'].startswith(alpha[c]):
                print "    |- {fid}: {name}".format(fid=f, name=findings[f]['name'])
                count += 1
                #if severity == 'Critical': # downgrade
                #    document.add_paragraph("{} (High Risk)".format(findings[f]['name']), style='High')
                #else:
                document.add_paragraph("{} ({} Risk)".format(findings[f]['name'], severity), style=severity)
                #title = document.add_paragraph('', "Heading 3")
                #title_run = title.add_run()
                #title_run.font.name = "Arial"
                #title_run.font.size = Pt(12)
                #title_run.font.bold = True
                #title_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                #title_run.add_text(findings[f]['name'] + " (" + findings[f]['severity'] + " Risk)")
                # HOSTLIST
                try:
                    syst = document.add_paragraph(style='framed')
                    syst.add_run('Impacted Systems:\n').bold = True
                    u_hosts = set(findings[f]['hosts'])
                    unique_hosts = sorted(u_hosts)
                    syst.add_run( ', '.join([str(x).strip() for x in unique_hosts]) ).italic = True
                except Exception as failedhostlist:
                    #print "Failed at: %s" % (failedhostlist)
                    pass
                # CVSS3 SCORE
                try:
                    if len(findings[f]['cvss3_base_score']) != 0:
                        cvss = document.add_paragraph(style='framed')
                        cvss.add_run('\nCVSS3 Score:\n').bold = True
                        cvss.add_run( findings[f]['cvss3_base_score'].strip()).italic = True
                    if len(findings[f]['cvss3_base_score']) != 0:
                        cvss = document.add_paragraph(style='framed')
                        cvss.add_run('\nCVSS3 Vector:\n').bold = True
                        cvss.add_run( findings[f]['cvss3_vector'].strip()).italic = True
                except Exception as failedcvvs:
                    #print "Failed at: %s" % (failedcvvs)
                    pass
                # BUGID
                try:
                    if len(findings[f]['bugid']) != 0:
                        #print "Len: %s" % (len(findings[f]['bugid']))
                        bigi = document.add_paragraph(style='framed')
                        u_bugids = set(findings[f]['bugid'])
                        unique_bugids = sorted(u_bugids)
                        bigi.add_run('\nBug ID:\n').bold = True
                        bigi.add_run(', '.join([str(x).strip() for x in unique_bugids]),).italic = True
                except Exception as failedbugid:
                    #print "Failed at: %s" % (failedbugid)
                    pass
                # DESCRIPTION
                try:
                    desc = document.add_paragraph(style='framed')
                    desc.add_run('\nDescription:\n').bold = True
                    desc.add_run( findings[f]['description'].strip())
                except Exception as faileddesc:
                    #print "Failed at: %s" % (faileddesc)
                    pass
                # CONSOLE OUTPUT
                try:
                    if findings[f]['output']is not None:
                        conout = document.add_paragraph(style='framed')
                        conout.add_run("\nCaptured Output:\n").bold = True
                        conout.add_run( findings[f]['output'].strip())
                except Exception as faileconout:
                    #print "Failed at: %s" % (faileconout)
                    pass
                # REMEDIATION
                try:					
                    recp = document.add_paragraph(style = 'framed')
                    recp.add_run('\nRemediation Recommendations:\n').bold = True
                    recp.add_run( findings[f]['recommendation'].strip())
                except Exception as failerecp:
                    #print "Failed at: %s" % (failerecp)
                    pass
                # EXTERNAL RESOURCES
                try:
                    if 'resources' in findings[f]:
                        reso = document.add_paragraph(style='framed')
                        reso.add_run('\nExternal Resources:\n').bold = True
                        reso.add_run( findings[f]['resources'].strip())
                except Exception as failedreso:
                    #print "Failed at: %s" % (failedreso)
                    pass
                # END OF FINDING CREATION
    print "    Total: {}".format(str(count))
'''
def printFindingsMD(findings, severity):
    for f in findings:
        if findings[f]['severity'] == severity:
            print "## " + findings[f]['name'] + " (" + findings[f]['severity'] + " Risk)"
            print "**Affected Systems:** " + ', '.join([str(x) for x in findings[f]['hosts']])
            print
            print "**Description:**"
            print findings[f]['description']
            print
            print "**Recommendation:**"
            print findings[f]['recommendation']
            print
            if 'resources' in findings[f]:
                print "**Resources:**"
                print findings[f]['resources']
                print
'''

def main(argv):
    filename = ""
    parser = argparse.ArgumentParser(description='Parse a XML dump.')
    parser.add_argument('inputfiles', type=file, help='the input  xml files. can be *.', nargs='+')
    #parser.add_argument('--hosts', action="store_true", help='Just print the discovered hosts')
    #parser.add_argument('--ports', action="store_true", help='Print discovered hosts and ports')
    #parser.add_argument('--delim', type=str, default=' ', required=False, help='delimiter to use for output.')

    try:
        args = parser.parse_args()
        parse( args )
    except IOError, err:
        print str(err)
        parser.print_help()
        sys.exit(2)

if __name__ == "__main__":
    main(sys.argv[1:])
